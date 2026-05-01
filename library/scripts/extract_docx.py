"""Extract structural blocks from a .docx (Word) file.

Word documents already carry the structure that the PDF pipeline has to
reverse-engineer with font-size heuristics: paragraphs are objects, heading
levels are styles, tables and footnotes are typed nodes. So this extractor
is a thin mapping layer — it walks paragraphs in document order, looks up
each one's style, and emits the same Block schema produced by extract.py.

DOCX has no notion of a printed-page anchor (page breaks are renderer
output, not source structure), so every block sets `print_page=""`. The
emitter in tex_emit.py already skips `\\pgmark{}` for empty print_page,
so DOCX papers render cleanly without page chips.
"""

from __future__ import annotations

from dataclasses import asdict
from typing import Any

from extract import Block


_HEADING_PREFIX = "heading "  # python-docx style names: "Heading 1", "Heading 2", ...


def _classify_paragraph(style_name: str, text: str) -> tuple[str, int]:
    """Return (block_kind, heading_level). heading_level is 0 unless kind=heading."""
    name = (style_name or "").strip().lower()
    if name.startswith(_HEADING_PREFIX):
        try:
            level = int(name[len(_HEADING_PREFIX):].strip() or "1")
        except ValueError:
            level = 1
        return "heading", max(1, min(level, 3))
    if name == "title":
        return "heading", 1
    if name == "subtitle":
        return "heading", 2
    if name == "caption":
        return "caption", 0
    if name in ("footnote text", "endnote text"):
        return "footnote", 0
    if name in ("list paragraph", "list bullet", "list number"):
        return "list_item", 0
    return "text", 0


def _extract_docx(docx_path: str) -> list[Block]:
    from docx import Document

    doc = Document(docx_path)
    blocks: list[Block] = []

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style_name = para.style.name if para.style is not None else ""
        kind, level = _classify_paragraph(style_name, text)
        blocks.append(Block(kind=kind, level=level, text=text, pdf_page=0, print_page=""))

    # Tables: flatten row by row. Each row becomes a single text block so it
    # at least round-trips into the rendered paper. A future iteration could
    # emit a real `tabular` block kind.
    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip().replace("\n", " ") for cell in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                blocks.append(Block(kind="text", text=row_text, pdf_page=0, print_page=""))

    # Footnotes live in a separate part. python-docx exposes them via the
    # underlying lxml; we read text only and append them at the end as
    # standalone footnote blocks. They won't be re-attached to bodies (same
    # caveat as the PDF marker fallback in extract.py).
    try:
        footnotes_part = doc.part.package.part_related_by_type(
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
        )
    except Exception:
        footnotes_part = None
    if footnotes_part is not None:
        try:
            from docx.oxml.ns import qn  # type: ignore

            ns_w_t = qn("w:t")
            ns_w_footnote = qn("w:footnote")
            root = footnotes_part.element
            for fn in root.iter(ns_w_footnote):
                # Skip the separator / continuation pseudo-footnotes (id 0, -1).
                fn_id = fn.get(qn("w:id"))
                if fn_id in (None, "0", "-1"):
                    continue
                texts = [t.text or "" for t in fn.iter(ns_w_t)]
                joined = "".join(texts).strip()
                if joined:
                    blocks.append(
                        Block(kind="footnote", text=joined, pdf_page=0, print_page="")
                    )
        except Exception:
            pass

    return blocks


def extract(docx_path: str) -> tuple[list[Block], str]:
    """Match the public surface of extract.extract(). Returns (blocks, name)."""
    return _extract_docx(docx_path), "docx-native"


def extract_to_json(docx_path: str) -> dict:
    blocks, extractor = extract(docx_path)
    return {
        "extractor": extractor,
        "blocks": [asdict(b) for b in blocks],
    }


def core_properties(docx_path: str) -> dict[str, Any]:
    """Read DOCX core properties for triage (title, author, etc.)."""
    from docx import Document

    doc = Document(docx_path)
    cp = doc.core_properties
    first_paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            first_paragraphs.append(text)
        if len(first_paragraphs) >= 8:
            break
    return {
        "title": cp.title or "",
        "author": cp.author or "",
        "created": cp.created.isoformat() if cp.created else "",
        "modified": cp.modified.isoformat() if cp.modified else "",
        "first_paragraphs": first_paragraphs,
    }


if __name__ == "__main__":
    import json
    import sys

    if len(sys.argv) < 2:
        print("usage: python extract_docx.py <docx>")
        sys.exit(2)
    print(json.dumps(extract_to_json(sys.argv[1]), indent=2))
